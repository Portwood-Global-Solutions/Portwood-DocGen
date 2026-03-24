from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_donation_receipt_with_logo():
    doc = Document()

    # Global Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)

    # --- BRANDED LOGO HEADER (using a table for the 'Logo' look) ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(6.5)
    
    # Left Cell: The 'Logo'
    logo_cell = header_table.rows[0].cells[0]
    shading_elm = parse_xml(r'<w:shd {} w:fill="2C3E50"/>'.format(nsdecls('w')))
    logo_cell._tc.get_or_add_tcPr().append(shading_elm)
    
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_p.add_run("GP") # Logo Initials
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(255, 255, 255) # White

    # Right Cell: Org Info
    info_cell = header_table.rows[0].cells[1]
    info_p = info_cell.paragraphs[0]
    info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = info_p.add_run("THE GENEROSITY PROJECT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    info_p.add_run("\n123 Mission Way, Suite 100\nSan Francisco, CA 94105\nTax ID: 12-3456789")
    
    doc.add_paragraph().paragraph_format.space_after = Inches(0.4)

    # --- DOCUMENT TITLE ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OFFICIAL DONATION RECEIPT")
    run.bold = True
    run.font.size = Pt(16)
    run.underline = True

    doc.add_paragraph().paragraph_format.space_after = Inches(0.2)

    # --- DONOR ADDRESS BLOCK ---
    donor_block = doc.add_paragraph()
    donor_block.add_run("{#OpportunityContactRoles}")
    donor_block.add_run("{Contact.FirstName} {Contact.LastName}\n")
    donor_block.add_run("{Contact.MailingStreet}\n")
    donor_block.add_run("{Contact.MailingCity}, {Contact.MailingState} {Contact.MailingPostalCode}\n")
    donor_block.add_run("{Contact.MailingCountry}")
    donor_block.add_run("{/OpportunityContactRoles}")
    donor_block.paragraph_format.space_after = Inches(0.3)

    # --- RECEIPT DETAILS ---
    details = doc.add_paragraph()
    details.add_run("Receipt Number: ").bold = True
    details.add_run("RCPT-{Id}\n")
    details.add_run("Date of Gift: ").bold = True
    details.add_run("March 24, 2026\n")
    details.add_run("Donation Amount: ").bold = True
    details.add_run("{Amount:currency}")
    
    doc.add_paragraph().paragraph_format.space_after = Inches(0.3)

    # --- BODY ---
    salutation = doc.add_paragraph()
    salutation.add_run("Dear {#OpportunityContactRoles}{Contact.FirstName}{/OpportunityContactRoles},")

    body = doc.add_paragraph()
    body.add_run("\nOn behalf of everyone at ")
    body.add_run("{Account.Name}").bold = True
    body.add_run(", we would like to express our deepest thanks for your donation to the ")
    body.add_run("{Name}").italic = True
    body.add_run(" fund.")

    statement = doc.add_paragraph()
    statement.add_run("\nThis letter serves as your official receipt for tax purposes. No goods or services were provided in exchange for this contribution.")
    statement.font.size = Pt(10)
    statement.italic = True
    
    doc.add_paragraph().paragraph_format.space_after = Inches(0.5)

    # --- SIGNATURE ---
    sig = doc.add_paragraph()
    sig.add_run("Authorized Signature:\n")
    sig.add_run("___________________________\n")
    sig.add_run("Executive Director").bold = True

    # Save
    doc.save("Donation_Receipt_Template.docx")
    print("Donation_Receipt_Template.docx updated with branded logo header.")

if __name__ == "__main__":
    create_donation_receipt_with_logo()
